"""Document normalization, hashing, and chunking."""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path

_WHITESPACE_RE = re.compile(r"[ \t\f\v]+")
_PARAGRAPH_RE = re.compile(r"\n{2,}")
_PUBLIC_ID_RE = re.compile(r"^[a-z0-9][a-z0-9_.:-]{0,63}$")
_PRIVATE_ID_KEYWORDS = frozenset(
    {
        "assignment",
        "draft",
        "essay",
        "homework",
        "learner",
        "name",
        "paper",
        "person",
        "student",
        "submission",
        "teacher",
        "user",
    }
)
_DOCUMENT_EXTENSION_RE = re.compile(r"\.(docx?|pdf|md|markdown|txt|rtf|odt)\b", re.IGNORECASE)


@dataclass(frozen=True)
class TextSegment:
    """A normalized document segment with character offsets."""

    kind: str
    index: int
    start: int
    end: int
    label: str | None = None


@dataclass(frozen=True)
class Document:
    """A normalized document."""

    id: str
    text: str
    sha256: str
    metadata: dict[str, str]
    segments: tuple[TextSegment, ...] = ()

    @classmethod
    def from_text(cls, text: str, *, document_id: str | None = None, **metadata: str) -> Document:
        normalized = normalize_text(text)
        digest = sha256_text(normalized)
        safe_id = public_document_id(document_id, fallback_digest=digest)
        return cls(
            id=safe_id,
            text=normalized,
            sha256=digest,
            metadata=metadata,
            segments=tuple(_paragraph_segments(normalized)),
        )


def load_document_file(path: str | Path, *, document_id: str | None = None) -> Document:
    """Load a supported local document into normalized text."""

    document_path = Path(path)
    suffix = document_path.suffix.casefold()
    if suffix in {".txt", ".md", ".markdown"}:
        text = document_path.read_text(encoding="utf-8", errors="ignore")
        return Document.from_text(text, document_id=document_id, file_type=suffix.lstrip("."))
    elif suffix == ".pdf":
        return _document_from_segment_texts(
            _read_pdf_pages(document_path),
            document_id=document_id,
            segment_kind="page",
            file_type="pdf",
        )
    elif suffix == ".docx":
        return _document_from_segment_texts(
            _read_docx_paragraphs(document_path),
            document_id=document_id,
            segment_kind="paragraph",
            file_type="docx",
        )
    else:
        raise ValueError(f"Unsupported document type: {suffix or 'none'}")


@dataclass(frozen=True)
class TextChunk:
    """A chunk of a normalized document."""

    id: str
    document_id: str
    text: str
    start: int
    end: int


def normalize_text(text: str) -> str:
    """Normalize whitespace while preserving paragraph boundaries."""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [_WHITESPACE_RE.sub(" ", line).strip() for line in text.split("\n")]
    normalized = "\n".join(lines)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def sha256_text(text: str) -> str:
    """Return the SHA-256 digest for UTF-8 text."""

    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def public_document_id(
    value: str | None,
    *,
    fallback_digest: str | None = None,
    prefix: str = "doc",
) -> str:
    """Return a report-safe document ID.

    This keeps short opaque lowercase IDs such as ``source-a`` but avoids exporting
    local paths, emails, filenames, title-like IDs, or other user-controlled
    provenance.
    """

    if value and _PUBLIC_ID_RE.fullmatch(value) and not _looks_like_private_identifier(value):
        return value
    digest = fallback_digest or sha256_text(value or prefix)
    return f"{prefix}:{digest[:16]}"


def is_public_document_id(value: object) -> bool:
    """Return whether a value can be safely used as a public document label."""

    return isinstance(value, str) and bool(value) and public_document_id(value) == value


def _looks_like_private_identifier(value: str) -> bool:
    lowered = value.casefold()
    parts = re.split(r"[-_.:]+", lowered)
    return any(
        (
            "@" in value,
            "/" in value,
            "\\" in value,
            value != lowered,
            lowered.startswith(("c:", "file:")),
            ".." in value,
            _DOCUMENT_EXTENSION_RE.search(value) is not None,
            any(part in _PRIVATE_ID_KEYWORDS for part in parts),
        )
    )


def _read_pdf_pages(path: Path) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "PDF ingestion requires the documents extra: python -m pip install -e \".[documents]\""
        ) from exc
    reader = PdfReader(str(path))
    return [page.extract_text() or "" for page in reader.pages]


def _read_docx_paragraphs(path: Path) -> list[str]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError(
            "DOCX ingestion requires the documents extra: python -m pip install -e \".[documents]\""
        ) from exc
    document = DocxDocument(str(path))
    return [paragraph.text for paragraph in document.paragraphs]


def _document_from_segment_texts(
    segment_texts: list[str],
    *,
    document_id: str | None,
    segment_kind: str,
    file_type: str,
) -> Document:
    normalized_segments = [normalize_text(text) for text in segment_texts]
    normalized_segments = [text for text in normalized_segments if text]
    normalized = "\n\n".join(normalized_segments)
    digest = sha256_text(normalized)
    safe_id = public_document_id(document_id, fallback_digest=digest)
    segments: list[TextSegment] = []
    cursor = 0
    for index, text in enumerate(normalized_segments):
        start = cursor
        end = start + len(text)
        segments.append(
            TextSegment(
                kind=segment_kind,
                index=index,
                start=start,
                end=end,
                label=f"{segment_kind} {index + 1}",
            )
        )
        cursor = end + 2
    return Document(
        id=safe_id,
        text=normalized,
        sha256=digest,
        metadata={"file_type": file_type},
        segments=tuple(segments),
    )


def _paragraph_segments(text: str) -> list[TextSegment]:
    return [
        TextSegment(
            kind="paragraph",
            index=index,
            start=start,
            end=start + len(paragraph),
            label=f"paragraph {index + 1}",
        )
        for index, (paragraph, start) in enumerate(_iter_paragraphs(text))
    ]


def chunk_text(
    text: str,
    *,
    document_id: str = "document",
    max_chars: int = 900,
    overlap_chars: int = 120,
) -> list[TextChunk]:
    """Split text into paragraph-aware chunks with optional overlap."""

    if max_chars <= 0:
        raise ValueError("max_chars must be positive")
    if overlap_chars < 0:
        raise ValueError("overlap_chars must be non-negative")
    if overlap_chars >= max_chars:
        raise ValueError("overlap_chars must be smaller than max_chars")

    normalized = normalize_text(text)
    if not normalized:
        return []

    chunks: list[TextChunk] = []
    for paragraph, paragraph_start in _iter_paragraphs(normalized):
        if len(paragraph) <= max_chars:
            chunks.append(
                _chunk(
                    document_id,
                    paragraph,
                    paragraph_start,
                    paragraph_start + len(paragraph),
                )
            )
            continue

        step = max_chars - overlap_chars
        offset = 0
        while offset < len(paragraph):
            end_offset = min(offset + max_chars, len(paragraph))
            chunk_text_value = paragraph[offset:end_offset].strip()
            if chunk_text_value:
                window = paragraph[offset:end_offset]
                start = paragraph_start + offset + (len(window) - len(window.lstrip()))
                end = start + len(chunk_text_value)
                chunks.append(_chunk(document_id, chunk_text_value, start, end))
            if end_offset == len(paragraph):
                break
            offset += step

    return [
        TextChunk(
            id=f"{document_id}:{index}",
            document_id=document_id,
            text=item.text,
            start=item.start,
            end=item.end,
        )
        for index, item in enumerate(chunks)
    ]


def _iter_paragraphs(text: str) -> list[tuple[str, int]]:
    paragraphs: list[tuple[str, int]] = []
    cursor = 0
    for match in _PARAGRAPH_RE.finditer(text):
        paragraph = text[cursor : match.start()].strip()
        if paragraph:
            window = text[cursor : match.start()]
            start = cursor + len(window) - len(window.lstrip())
            paragraphs.append((paragraph, start))
        cursor = match.end()
    tail = text[cursor:].strip()
    if tail:
        start = cursor + len(text[cursor:]) - len(text[cursor:].lstrip())
        paragraphs.append((tail, start))
    return paragraphs


def _chunk(document_id: str, text: str, start: int, end: int) -> TextChunk:
    return TextChunk(
        id=f"{document_id}:pending",
        document_id=document_id,
        text=text,
        start=start,
        end=end,
    )
